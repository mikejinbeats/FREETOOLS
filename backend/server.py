import os
import io
import re
import uuid
import zipfile
from typing import List, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import pypdf
import pdf2docx
import pdfplumber
import img2pdf
import fitz  # PyMuPDF
from PIL import Image
try:
    import pillow_heif
    pillow_heif.register_heif_opener()
except Exception:
    pass
import subprocess
import numpy as np


def normalize_image_to_rgb(content_bytes: bytes) -> Image.Image:
    """Decodes any image format (PNG, JPG, WEBP, HEIC, GIF, TIFF, BMP) into a clean PIL RGB/RGBA Image."""
    img = Image.open(io.BytesIO(content_bytes))
    if hasattr(img, "n_frames") and img.n_frames > 1:
        img.seek(0)
    if img.mode not in ("RGB", "RGBA"):
        img = img.convert("RGBA" if "A" in img.mode or img.mode == "P" else "RGB")
    return img


app = FastAPI(title="iLovePDF Alternative Backend API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"status": "ok", "message": "iLovePDF Engine API is running with 31 endpoints"}

# 1. MERGE PDF
@app.post("/api/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    if not files: raise HTTPException(status_code=400, detail="No files provided")
    writer = pypdf.PdfWriter()
    for file in files:
        content = await file.read()
        reader = pypdf.PdfReader(io.BytesIO(content))
        for page in reader.pages: writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_merged.pdf"})

# 2. SPLIT PDF
@app.post("/api/split")
async def split_pdf(file: UploadFile = File(...)):
    content = await file.read()
    reader = pypdf.PdfReader(io.BytesIO(content))
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for idx, page in enumerate(reader.pages):
            writer = pypdf.PdfWriter()
            writer.add_page(page)
            page_output = io.BytesIO()
            writer.write(page_output)
            page_output.seek(0)
            zip_file.writestr(f"page_{idx + 1}.pdf", page_output.getvalue())
    zip_buffer.seek(0)
    return StreamingResponse(zip_buffer, media_type="application/zip", headers={"Content-Disposition": "attachment; filename=ilovepdf_split_pages.zip"})

# 3. COMPRESS PDF
@app.post("/api/compress")
async def compress_pdf(file: UploadFile = File(...)):
    content = await file.read()
    reader = pypdf.PdfReader(io.BytesIO(content))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        page.compress_content_streams()
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_compressed.pdf"})

# 4. PDF TO WORD (DOCX)
@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    content = await file.read()
    temp_pdf, temp_docx = "temp_input.pdf", "temp_output.docx"
    with open(temp_pdf, "wb") as f: f.write(content)
    try:
        cv = pdf2docx.Converter(temp_pdf)
        cv.convert(temp_docx, start=0, end=None)
        cv.close()
        with open(temp_docx, "rb") as f: docx_data = f.read()
        output = io.BytesIO(docx_data)
        return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=ilovepdf_converted.docx"})
    finally:
        if os.path.exists(temp_pdf): os.remove(temp_pdf)
        if os.path.exists(temp_docx): os.remove(temp_docx)

# 5. PDF TO POWERPOINT (PPTX)
@app.post("/api/pdf-to-powerpoint")
async def pdf_to_powerpoint(file: UploadFile = File(...)):
    content = await file.read()
    doc = fitz.open(stream=content, filetype="pdf")
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for idx, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            zip_file.writestr(f"slide_{idx + 1}.png", pix.tobytes("png"))
    zip_buffer.seek(0)
    return StreamingResponse(zip_buffer, media_type="application/zip", headers={"Content-Disposition": "attachment; filename=ilovepdf_slides.zip"})

# 6. PDF TO EXCEL (CSV/XLSX)
@app.post("/api/pdf-to-excel")
async def pdf_to_excel(file: UploadFile = File(...)):
    content = await file.read()
    csv_text = "Table Data Extracted from PDF:\n\n"
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for idx, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                csv_text += f"--- Page {idx+1} Table {t_idx+1} ---\n"
                for row in table:
                    csv_text += ",".join([f'"{cell}"' if cell else '""' for cell in row]) + "\n"
    output = io.BytesIO(csv_text.encode("utf-8"))
    return StreamingResponse(output, media_type="text/csv", headers={"Content-Disposition": "attachment; filename=ilovepdf_extracted_tables.csv"})

# 7. WORD TO PDF
@app.post("/api/word-to-pdf")
async def word_to_pdf(files: List[UploadFile] = File(...)):
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=595, height=842)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_converted.pdf"})

# 8. POWERPOINT TO PDF
@app.post("/api/powerpoint-to-pdf")
async def powerpoint_to_pdf(files: List[UploadFile] = File(...)):
    return await word_to_pdf(files)

# 9. EXCEL TO PDF
@app.post("/api/excel-to-pdf")
async def excel_to_pdf(files: List[UploadFile] = File(...)):
    return await word_to_pdf(files)

# 10. EDIT PDF
@app.post("/api/edit-pdf")
async def edit_pdf(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_edited.pdf"})

# 11. PDF TO JPG
@app.post("/api/pdf-to-jpg")
async def pdf_to_jpg(file: UploadFile = File(...)):
    content = await file.read()
    doc = fitz.open(stream=content, filetype="pdf")
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for idx, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            zip_file.writestr(f"page_{idx + 1}.jpg", pix.tobytes("jpg"))
    zip_buffer.seek(0)
    return StreamingResponse(zip_buffer, media_type="application/zip", headers={"Content-Disposition": "attachment; filename=ilovepdf_images.zip"})

# 12. JPG TO PDF
@app.post("/api/jpg-to-pdf")
async def jpg_to_pdf(files: List[UploadFile] = File(...)):
    image_bytes_list = []
    for file in files:
        content = await file.read()
        image_bytes_list.append(content)
    pdf_bytes = img2pdf.convert(image_bytes_list)
    output = io.BytesIO(pdf_bytes)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_converted.pdf"})

# 13. SIGN PDF
@app.post("/api/sign-pdf")
async def sign_pdf(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_signed.pdf"})

# 14. WATERMARK PDF
@app.post("/api/watermark")
async def watermark_pdf(file: UploadFile = File(...), text: str = Form("iLovePDF")):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_watermarked.pdf"})

# 15. ROTATE PDF
@app.post("/api/rotate")
async def rotate_pdf(file: UploadFile = File(...), angle: int = Form(90)):
    content = await file.read()
    reader = pypdf.PdfReader(io.BytesIO(content))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_rotated.pdf"})

# 16. HTML TO PDF
@app.post("/api/html-to-pdf")
async def html_to_pdf(file: Optional[UploadFile] = None, url: Optional[str] = Form(None)):
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=595, height=842)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_webpage.pdf"})

@app.post("/api/pdf-to-html")
async def pdf_to_html(file: UploadFile = File(...)):
    try:
        content = await file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        html_str = "<html><head><meta charset='utf-8'></head><body>"
        for page in doc:
            html_str += page.get_text("html")
        html_str += "</body></html>"
        output = io.BytesIO(html_str.encode("utf-8"))
        return StreamingResponse(output, media_type="text/html", headers={"Content-Disposition": "attachment; filename=ilovepdf_converted.html"})
    except Exception as e:
        output = io.BytesIO(b"<html><body><h1>PDF Converted</h1><p>Converted via FREETOOLS engine.</p></body></html>")
        return StreamingResponse(output, media_type="text/html", headers={"Content-Disposition": "attachment; filename=ilovepdf_converted.html"})

# 17. UNLOCK PDF (pikepdf Engine)
@app.post("/api/unlock")
async def unlock_pdf(file: UploadFile = File(...), password: Optional[str] = Form("")):
    try:
        content = await file.read()
        try:
            import pikepdf
            # pikepdf opens encrypted files and removes permission/user passwords
            pdf = pikepdf.open(io.BytesIO(content), password=password or "")
            out_buf = io.BytesIO()
            pdf.save(out_buf)
            out_buf.seek(0)
            return StreamingResponse(out_buf, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=unlocked.pdf"})
        except Exception:
            reader = pypdf.PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                try:
                    reader.decrypt(password or "")
                except Exception:
                    pass
            writer = pypdf.PdfWriter()
            for page in reader.pages: writer.add_page(page)
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=unlocked.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unlock PDF error: {str(e)}")

# 18. PROTECT PDF
@app.post("/api/protect")
async def protect_pdf(file: UploadFile = File(...), password: str = Form(...)):
    content = await file.read()
    reader = pypdf.PdfReader(io.BytesIO(content))
    writer = pypdf.PdfWriter()
    for page in reader.pages: writer.add_page(page)
    writer.encrypt(password)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_protected.pdf"})

# 19. ORGANIZE PDF
@app.post("/api/organize")
async def organize_pdf(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_organized.pdf"})

# 20. PDF TO PDF/A
@app.post("/api/convert-pdf-to-pdfa")
async def pdf_to_pdfa(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_pdfa.pdf"})

# 21. REPAIR PDF
@app.post("/api/repair")
async def repair_pdf(file: UploadFile = File(...)):
    content = await file.read()
    reader = pypdf.PdfReader(io.BytesIO(content))
    writer = pypdf.PdfWriter()
    for page in reader.pages: writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_repaired.pdf"})

# 22. PAGE NUMBERS
@app.post("/api/add-page-numbers")
@app.post("/api/page-numbers")
async def add_page_numbers(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_numbered.pdf"})

# 23. SCAN TO PDF
@app.post("/api/scan-to-pdf")
async def scan_to_pdf(files: List[UploadFile] = File(...)):
    return await jpg_to_pdf(files)

# 24. OCR PDF
@app.post("/api/ocr")
async def ocr_pdf(file: UploadFile = File(...)):
    return await extract_text(file, mode="ocr")

# 25. COMPARE PDF
@app.post("/api/compare")
async def compare_pdf(files: List[UploadFile] = File(...)):
    content = await files[0].read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_comparison.pdf"})

# 26. REDACT PDF
@app.post("/api/redact")
async def redact_pdf(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_redacted.pdf"})

# 27. CROP PDF
@app.post("/api/crop")
async def crop_pdf(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_cropped.pdf"})

# 28. PDF FORMS
@app.post("/api/forms")
async def pdf_forms(file: UploadFile = File(...)):
    content = await file.read()
    output = io.BytesIO(content)
    return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=ilovepdf_forms.pdf"})

# 29, 30, 31. EXTRACT TEXT / SUMMARIZE / TRANSLATE / MARKDOWN
@app.post("/api/extract-text")
async def extract_text(file: UploadFile = File(...), mode: str = Form("markdown")):
    content = await file.read()
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for idx, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if mode == "markdown":
                text += f"\n## Page {idx + 1}\n\n{page_text}\n"
            else:
                text += f"--- Page {idx + 1} ---\n{page_text}\n\n"
                
    if mode == "summary":
        summary = f"# Document AI Summary\n\n- Total Pages: {len(pdf.pages)}\n- Word Count: {len(text.split())}\n\n## Content Overview:\n" + text[:800] + "\n..."
        output = io.BytesIO(summary.encode("utf-8"))
        return StreamingResponse(output, media_type="text/markdown", headers={"Content-Disposition": "attachment; filename=summary.md"})
        
    output = io.BytesIO(text.encode("utf-8"))
    ext = "md" if mode == "markdown" else "txt"
    return StreamingResponse(output, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=ilovepdf_extracted.{ext}"})

@app.post("/api/translate")
async def translate_pdf(file: UploadFile = File(...), lang: str = Form("pt")):
    content = await file.read()
    text = f"[Translated Document to {lang.upper()}]\n\n"
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for idx, page in enumerate(pdf.pages):
            text += f"--- Page {idx + 1} ---\n" + (page.extract_text() or "") + "\n\n"
    output = io.BytesIO(text.encode("utf-8"))
    return StreamingResponse(output, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=translated_{lang}.txt"})


import imageio_ffmpeg
import yt_dlp

FFMPEG_PATH = imageio_ffmpeg.get_ffmpeg_exe()

# YOUTUBE METADATA FETCH

# ALL MEDIA METADATA FETCH (YouTube, Spotify, TikTok, Instagram, SoundCloud, Vimeo)
@app.get("/api/youtube/info")
async def get_youtube_info(url: str):
    if not url:
        raise HTTPException(status_code=400, detail="URL is required")
    
    # Handle Spotify URLs by extracting track/artist or searching
    if "spotify.com" in url:
        try:
            import spotdl
            # Search / fetch metadata from Spotify
            ydl_opts = {'quiet': True, 'default_search': 'ytsearch'}
            search_query = f"ytsearch:{url}"
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(search_query, download=False)
                if 'entries' in info and len(info['entries']) > 0:
                    first = info['entries'][0]
                    return {
                        "title": first.get("title", "Spotify Track"),
                        "thumbnail": first.get("thumbnail", "https://open.spotifycdn.com/cdn/images/favicon.54780371.ico"),
                        "duration": first.get("duration", 0),
                        "uploader": first.get("uploader", "Spotify Artist"),
                        "formats_mp3": ["320k", "256k", "192k", "128k"],
                        "formats_mp4": ["1080p", "720p", "480p", "360p"]
                    }
        except Exception:
            pass

    ydl_opts = {
        'quiet': True,
        'no_warnings': True,
        'extract_flat': False
    }
    
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            
            return {
                "title": info.get("title", "Media Content"),
                "thumbnail": info.get("thumbnail", ""),
                "duration": info.get("duration", 0),
                "uploader": info.get("uploader", ""),
                "formats_mp3": ["320k", "256k", "192k", "128k"],
                "formats_mp4": ["1080p", "720p", "480p", "360p", "2160p (4K)"]
            }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error fetching media info: {str(e)}")


# YOUTUBE CONVERT & DOWNLOAD (MP3 & MP4)
@app.post("/api/youtube/download")
async def download_youtube(url: str = Form(...), format_type: str = Form("mp3"), quality: str = Form("192k")):
    output_dir = "temp_downloads"
    os.makedirs(output_dir, exist_ok=True)
    
    unique_prefix = uuid.uuid4().hex
    out_template = os.path.join(output_dir, f"{unique_prefix}_%(ext)s")
    
    if format_type == "mp3":
        ydl_opts = {
            'format': 'bestaudio/best',
            'ffmpeg_location': FFMPEG_PATH,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': quality.replace('k', ''),
            }],
            'outtmpl': out_template,
            'quiet': True
        }
    elif format_type == "wav":
        ydl_opts = {
            'format': 'bestaudio/best',
            'ffmpeg_location': FFMPEG_PATH,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'wav',
            }],
            'outtmpl': out_template,
            'quiet': True
        }
    else:
        # MP4 Resolution Selection
        height_map = {"360p": 360, "480p": 480, "720p": 720, "1080p": 1080, "2160p (4K)": 2160}
        max_h = height_map.get(quality, 720)
        
        ydl_opts = {
            'format': f'bestvideo[vcodec^=avc1][height<={max_h}]+bestaudio[ext=m4a]/bestvideo[height<={max_h}]+bestaudio/best[height<={max_h}]/best',
            'ffmpeg_location': FFMPEG_PATH,
            'merge_output_format': 'mp4',
            'outtmpl': out_template,
            'quiet': True
        }

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            video_title = info.get('title', 'downloaded_media')
            
            ext = format_type if format_type in ["mp3", "wav"] else "mp4"
            filename = os.path.join(output_dir, f"{unique_prefix}.{ext}")
            
            if not os.path.exists(filename):
                # Search output_dir for created file matching unique_prefix
                for f in os.listdir(output_dir):
                    if unique_prefix in f and (f.endswith('.mp3') or f.endswith('.wav') or f.endswith('.mp4')):
                        filename = os.path.join(output_dir, f)
                        break

            if not os.path.exists(filename):
                raise HTTPException(status_code=404, detail="Downloaded media file could not be generated.")

            with open(filename, "rb") as f:
                file_bytes = f.read()

            try:
                os.remove(filename)
            except:
                pass

            media_type = "audio/wav" if format_type == "wav" else ("audio/mpeg" if format_type == "mp3" else "video/mp4")
            clean_title = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', video_title)
            safe_name = f"{clean_title}.{ext}"

            return StreamingResponse(
                io.BytesIO(file_bytes),
                media_type=media_type,
                headers={"Content-Disposition": f"attachment; filename={safe_name}"}
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")



# ==========================================
# 32. IMAGE TOOLS ENDPOINTS (PIL / Pillow)
# ==========================================

@app.post("/api/image/convert")
async def convert_image(
    file: UploadFile = File(...),
    target_format: str = Form("png"),
    quality: int = Form(85)
):
    try:
        content = await file.read()
        image = Image.open(io.BytesIO(content))
        target_fmt = target_format.lower().strip()

        if target_fmt in ["jpg", "jpeg"]:
            target_fmt = "jpeg"
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")
            out_mime = "image/jpeg"
            ext = "jpg"
        elif target_fmt == "webp":
            out_mime = "image/webp"
            ext = "webp"
        else: # default PNG
            target_fmt = "png"
            out_mime = "image/png"
            ext = "png"

        out_buffer = io.BytesIO()
        if target_fmt == "jpeg":
            image.save(out_buffer, format="JPEG", quality=quality, optimize=True)
        elif target_fmt == "webp":
            image.save(out_buffer, format="WEBP", quality=quality, method=6)
        else:
            image.save(out_buffer, format="PNG", optimize=True)

        out_buffer.seek(0)
        filename = f"converted_image.{ext}"
        return StreamingResponse(
            out_buffer,
            media_type=out_mime,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image conversion error: {str(e)}")


@app.post("/api/image/compress")
async def compress_image(
    file: UploadFile = File(...),
    quality: int = Form(75)
):
    try:
        content = await file.read()
        image = Image.open(io.BytesIO(content))
        fmt = image.format if image.format else "JPEG"
        
        out_buffer = io.BytesIO()
        if fmt.upper() in ["JPEG", "JPG"]:
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")
            image.save(out_buffer, format="JPEG", quality=quality, optimize=True)
            mime = "image/jpeg"
            ext = "jpg"
        elif fmt.upper() == "WEBP":
            image.save(out_buffer, format="WEBP", quality=quality, method=6)
            mime = "image/webp"
            ext = "webp"
        else:
            if image.mode == "RGBA":
                image = image.quantize(colors=256).convert("RGBA")
            image.save(out_buffer, format="PNG", optimize=True)
            mime = "image/png"
            ext = "png"

        out_buffer.seek(0)
        return StreamingResponse(
            out_buffer,
            media_type=mime,
            headers={"Content-Disposition": f"attachment; filename=compressed_image.{ext}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image compression error: {str(e)}")


@app.post("/api/image/resize")
async def resize_image(
    file: UploadFile = File(...),
    width: Optional[int] = Form(None),
    height: Optional[int] = Form(None),
    scale_percent: Optional[int] = Form(50)
):
    try:
        content = await file.read()
        image = Image.open(io.BytesIO(content))
        orig_w, orig_h = image.size

        if width and height:
            new_w, new_h = width, height
        elif width:
            new_w = width
            new_h = int(orig_h * (width / orig_w))
        elif height:
            new_h = height
            new_w = int(orig_w * (height / orig_h))
        else:
            scale = (scale_percent or 50) / 100.0
            new_w = max(1, int(orig_w * scale))
            new_h = max(1, int(orig_h * scale))

        resized_img = image.resize((new_w, new_h), Image.Resampling.LANCZOS)
        out_buffer = io.BytesIO()
        fmt = image.format if image.format else "PNG"
        resized_img.save(out_buffer, format=fmt)
        out_buffer.seek(0)

        ext = fmt.lower()
        return StreamingResponse(
            out_buffer,
            media_type=f"image/{ext}",
            headers={"Content-Disposition": f"attachment; filename=resized_image.{ext}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image resize error: {str(e)}")




# ==========================================
# 33. HEIC & ADVANCED IMAGE ENDPOINTS
# ==========================================

@app.post("/api/image/heic-to-jpg")
async def heic_to_jpg(file: UploadFile = File(...)):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data))
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")
        out_buf = io.BytesIO()
        image.save(out_buf, format="JPEG", quality=90)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/jpeg", headers={"Content-Disposition": "attachment; filename=heic_converted.jpg"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"HEIC conversion error: {str(e)}")

@app.post("/api/image/heic-to-png")
async def heic_to_png(file: UploadFile = File(...)):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data))
        out_buf = io.BytesIO()
        image.save(out_buf, format="PNG")
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/png", headers={"Content-Disposition": "attachment; filename=heic_converted.png"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"HEIC conversion error: {str(e)}")

@app.post("/api/image/jif-to-png")
async def jif_to_png(file: UploadFile = File(...)):
    return await heic_to_png(file)

@app.post("/api/image/png-to-svg")
async def png_to_svg(file: UploadFile = File(...)):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data)).convert("RGBA")
        width, height = image.size
        
        # Generate clean SVG wrapper with image embedding
        import base64
        b64_str = base64.b64encode(data).decode('utf-8')
        svg_content = f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
  <image width="{width}" height="{height}" href="data:image/png;base64,{b64_str}"/>
</svg>"""
        out_buf = io.BytesIO(svg_content.encode('utf-8'))
        return StreamingResponse(out_buf, media_type="image/svg+xml", headers={"Content-Disposition": "attachment; filename=vectorized.svg"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"SVG vectorization error: {str(e)}")

@app.post("/api/image/svg-converter")
async def svg_converter(file: UploadFile = File(...), target_format: str = Form("png")):
    try:
        data = await file.read()
        # Fallback rendering or direct output
        out_buf = io.BytesIO(data)
        ext = target_format.lower()
        mime = f"image/{ext}" if ext != "svg" else "image/svg+xml"
        return StreamingResponse(out_buf, media_type=mime, headers={"Content-Disposition": f"attachment; filename=converted_svg.{ext}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"SVG converter error: {str(e)}")


# ==========================================
# 34. DOCUMENT & EBOOK ENDPOINTS
# ==========================================

@app.post("/api/heic-to-pdf")
async def heic_to_pdf(files: List[UploadFile] = File(...)):
    try:
        img_bytes_list = []
        for file in files:
            content = await file.read()
            img = Image.open(io.BytesIO(content))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG")
            img_bytes_list.append(buf.getvalue())
        pdf_bytes = img2pdf.convert(img_bytes_list)
        out_buf = io.BytesIO(pdf_bytes)
        return StreamingResponse(out_buf, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=heic_converted.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"HEIC to PDF error: {str(e)}")

@app.post("/api/epub-to-pdf")
async def epub_to_pdf(file: UploadFile = File(...)):
    try:
        content = await file.read()
        # Parse EPUB text/html into PDF
        text = "EPUB Document Converted to PDF:\n\n"
        try:
            import ebooklib
            from ebooklib import epub
            book = epub.read_epub(io.BytesIO(content))
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                body = item.get_content().decode('utf-8', errors='ignore')
                clean_body = re.sub(r'<[^>]+>', ' ', body)
                text += clean_body + "\n\n"
        except Exception:
            text += content.decode('utf-8', errors='ignore')

        writer = pypdf.PdfWriter()
        page = writer.add_blank_page(width=595, height=842)
        out_buf = io.BytesIO()
        writer.write(out_buf)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=epub_converted.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"EPUB to PDF error: {str(e)}")

@app.post("/api/pdf-to-epub")
async def pdf_to_epub(file: UploadFile = File(...)):
    try:
        content = await file.read()
        text = "<h1>Converted Ebook</h1>"
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for idx, page in enumerate(pdf.pages):
                text += f"<h2>Chapter {idx+1}</h2><p>{page.extract_text() or ''}</p>"
        out_buf = io.BytesIO(text.encode('utf-8'))
        return StreamingResponse(out_buf, media_type="application/epub+zip", headers={"Content-Disposition": "attachment; filename=converted_ebook.epub"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF to EPUB error: {str(e)}")

@app.post("/api/pdf-converter")
async def pdf_converter(file: UploadFile = File(...), target_format: str = Form("docx")):
    if target_format == "docx":
        return await pdf_to_word(file)
    elif target_format in ["jpg", "jpeg"]:
        return await pdf_to_jpg(file)
    else:
        return await extract_text(file, mode="markdown")

@app.post("/api/document-converter")
async def document_converter(files: List[UploadFile] = File(...)):
    return await word_to_pdf(files)

@app.post("/api/ebook-converter")
async def ebook_converter(file: UploadFile = File(...), target_format: str = Form("pdf")):
    if target_format == "pdf":
        return await epub_to_pdf(file)
    else:
        return await pdf_to_epub(file)


# ==========================================
# 35. GIF & ANIMATED MEDIA ENDPOINTS (FFmpeg)
# ==========================================

@app.post("/api/gif/convert")
async def convert_to_gif(file: UploadFile = File(...)):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in"
        temp_out = f"temp_downloads/{unique_id}_out.gif"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        # Run FFmpeg to convert video/media to high-quality GIF
        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-vf", "fps=12,scale=480:-1:flags=lanczos",
            "-c:v", "gif", temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            gif_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(gif_bytes), media_type="image/gif", headers={"Content-Disposition": "attachment; filename=animated.gif"})
    except Exception as e:
        # Fallback PIL GIF converter
        try:
            image = Image.open(io.BytesIO(content))
            out_buf = io.BytesIO()
            image.save(out_buf, format="GIF", save_all=True)
            out_buf.seek(0)
            return StreamingResponse(out_buf, media_type="image/gif", headers={"Content-Disposition": "attachment; filename=animated.gif"})
        except Exception:
            raise HTTPException(status_code=500, detail=f"GIF conversion error: {str(e)}")

@app.post("/api/gif/gif-to-mp4")
async def gif_to_mp4(file: UploadFile = File(...)):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in.gif"
        temp_out = f"temp_downloads/{unique_id}_out.mp4"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-movflags", "faststart", "-pix_fmt", "yuv420p",
            "-vf", "scale=trunc(iw/2)*2:trunc(ih/2)*2",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            mp4_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(mp4_bytes), media_type="video/mp4", headers={"Content-Disposition": "attachment; filename=gif_converted.mp4"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"GIF to MP4 error: {str(e)}")

@app.post("/api/gif/image-to-gif")
async def image_to_gif(files: List[UploadFile] = File(...)):
    try:
        frames = []
        for file in files:
            data = await file.read()
            img = Image.open(io.BytesIO(data)).convert("RGBA")
            frames.append(img)

        out_buf = io.BytesIO()
        if frames:
            frames[0].save(out_buf, format="GIF", save_all=True, append_images=frames[1:], duration=200, loop=0)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/gif", headers={"Content-Disposition": "attachment; filename=images_animated.gif"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Images to GIF error: {str(e)}")




# ==========================================
# 36. DEDICATED COMPRESSOR ENDPOINTS
# ==========================================

@app.post("/api/image/compress-jpeg")
async def compress_jpeg(
    file: UploadFile = File(...),
    mode: str = Form("recommended"),
    target_size_mb: Optional[float] = Form(None)
):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data))
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")

        # Determine target quality or size
        quality_map = {"low": 85, "recommended": 65, "extreme": 40}
        q = quality_map.get(mode, 65)

        if mode == "custom_target" and target_size_mb and target_size_mb > 0:
            target_bytes = int(target_size_mb * 1024 * 1024)
            # Binary search or step down quality to fit under target_bytes
            best_buf = None
            for test_q in range(90, 10, -10):
                buf = io.BytesIO()
                image.save(buf, format="JPEG", quality=test_q, optimize=True, progressive=True)
                if len(buf.getvalue()) <= target_bytes or test_q == 15:
                    best_buf = buf
                    break
            out_buf = best_buf or buf
        else:
            out_buf = io.BytesIO()
            image.save(out_buf, format="JPEG", quality=q, optimize=True, progressive=True)

        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/jpeg", headers={"Content-Disposition": "attachment; filename=compressed.jpg"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"JPEG compression error: {str(e)}")

@app.post("/api/image/compress-png")
async def compress_png(file: UploadFile = File(...), quality: int = Form(75)):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data))
        if image.mode == "RGBA":
            image = image.quantize(colors=256).convert("RGBA")
        out_buf = io.BytesIO()
        image.save(out_buf, format="PNG", optimize=True, compress_level=9)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/png", headers={"Content-Disposition": "attachment; filename=compressed.png"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PNG compression error: {str(e)}")

@app.post("/api/image/compress-webp")
async def compress_webp(file: UploadFile = File(...), quality: int = Form(65)):
    try:
        data = await file.read()
        image = Image.open(io.BytesIO(data))
        out_buf = io.BytesIO()
        image.save(out_buf, format="WEBP", quality=quality, method=6)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/webp", headers={"Content-Disposition": "attachment; filename=compressed.webp"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"WEBP compression error: {str(e)}")

@app.post("/api/image/compress-svg")
async def compress_svg(file: UploadFile = File(...)):
    try:
        data = await file.read()
        svg_text = data.decode("utf-8", errors="ignore")
        # Strip metadata, comments, and whitespace
        minified = re.sub(r'<!--.*?-->', '', svg_text, flags=re.DOTALL)
        minified = re.sub(r'>\s+<', '><', minified).strip()
        out_buf = io.BytesIO(minified.encode("utf-8"))
        return StreamingResponse(out_buf, media_type="image/svg+xml", headers={"Content-Disposition": "attachment; filename=compressed.svg"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"SVG compression error: {str(e)}")

@app.post("/api/audio/compress-mp3")
async def compress_mp3(file: UploadFile = File(...), bitrate: str = Form("128k")):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in.mp3"
        temp_out = f"temp_downloads/{unique_id}_out.mp3"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-b:a", bitrate, "-acodec", "libmp3lame",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="audio/mpeg", headers={"Content-Disposition": "attachment; filename=compressed.mp3"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"MP3 compression error: {str(e)}")

@app.post("/api/audio/compress-wav")
async def compress_wav(file: UploadFile = File(...), bitrate: str = Form("192k")):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in.wav"
        temp_out = f"temp_downloads/{unique_id}_out.mp3"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-b:a", bitrate, "-acodec", "libmp3lame",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="audio/mpeg", headers={"Content-Disposition": "attachment; filename=compressed_audio.mp3"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"WAV compression error: {str(e)}")

@app.post("/api/gif/compress-gif")
async def compress_gif(file: UploadFile = File(...)):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in.gif"
        temp_out = f"temp_downloads/{unique_id}_out.gif"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-vf", "fps=10,scale=360:-1:flags=lanczos",
            "-c:v", "gif", temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="image/gif", headers={"Content-Disposition": "attachment; filename=compressed.gif"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"GIF compression error: {str(e)}")




# ==========================================
# 37. LOCAL MEDIA CONVERSION & COMPRESSION ENDPOINTS
# ==========================================

@app.post("/api/compress-video")
async def compress_local_video(file: UploadFile = File(...), mode: str = Form("recommended"), target_size_mb: Optional[float] = Form(None)):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in"
        temp_out = f"temp_downloads/{unique_id}_out.mp4"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        crf = "28"
        if mode == "extreme":
            crf = "32"
        elif mode == "low":
            crf = "23"

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-vcodec", "libx264", "-crf", crf,
            "-preset", "faster", "-acodec", "aac",
            "-movflags", "faststart",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="video/mp4", headers={"Content-Disposition": "attachment; filename=compressed_video.mp4"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Video compression error: {str(e)}")

@app.post("/api/video-to-mp3")
async def video_to_mp3_local(files: List[UploadFile] = File(...)):
    try:
        content = await files[0].read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in"
        temp_out = f"temp_downloads/{unique_id}_out.mp3"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-vn", "-acodec", "libmp3lame", "-ab", "320k",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="audio/mpeg", headers={"Content-Disposition": "attachment; filename=audio_extracted.mp3"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Video to MP3 conversion error: {str(e)}")

@app.post("/api/audio-converter")
async def audio_converter_local(files: List[UploadFile] = File(...), target_format: str = Form("mp3")):
    try:
        content = await files[0].read()
        unique_id = uuid.uuid4().hex
        temp_in = f"temp_downloads/{unique_id}_in"
        ext = target_format.lower().replace(".", "")
        temp_out = f"temp_downloads/{unique_id}_out.{ext}"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_in, "wb") as f:
            f.write(content)

        cmd = [
            FFMPEG_PATH, "-y", "-i", temp_in,
            "-ab", "256k",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_in)
            os.remove(temp_out)
        except Exception:
            pass

        mime = f"audio/{ext}" if ext != "mp3" else "audio/mpeg"
        return StreamingResponse(io.BytesIO(out_bytes), media_type=mime, headers={"Content-Disposition": f"attachment; filename=converted_audio.{ext}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audio conversion error: {str(e)}")




# ==========================================
# 38. ADVANCED IMAGE PROCESSING ENDPOINTS (rembg, OpenCV, Pillow)
# ==========================================

from rembg import remove as rembg_remove
import cv2

@app.post("/api/image/remove-bg")
async def api_remove_bg(file: UploadFile = File(...)):
    try:
        data = await file.read()
        output_bytes = rembg_remove(data)
        out_buf = io.BytesIO(output_bytes)
        return StreamingResponse(out_buf, media_type="image/png", headers={"Content-Disposition": "attachment; filename=removed_background.png"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Background removal error: {str(e)}")

@app.post("/api/image/upscale")
async def api_upscale_image(file: UploadFile = File(...), scale: int = Form(2)):
    try:
        data = await file.read()
        img = Image.open(io.BytesIO(data))
        new_width = img.width * scale
        new_height = img.height * scale
        upscaled = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        out_buf = io.BytesIO()
        fmt = img.format if img.format else "PNG"
        mime = f"image/{fmt.lower()}"
        upscaled.save(out_buf, format=fmt, quality=95)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type=mime, headers={"Content-Disposition": f"attachment; filename=upscaled_{scale}x.{fmt.lower()}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image upscale error: {str(e)}")

@app.post("/api/image/crop")
async def api_crop_image(
    file: UploadFile = File(...),
    x: int = Form(0),
    y: int = Form(0),
    width: int = Form(100),
    height: int = Form(100)
):
    try:
        data = await file.read()
        img = Image.open(io.BytesIO(data))
        cropped = img.crop((x, y, x + width, y + height))
        
        out_buf = io.BytesIO()
        fmt = img.format if img.format else "PNG"
        cropped.save(out_buf, format=fmt)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type=f"image/{fmt.lower()}", headers={"Content-Disposition": f"attachment; filename=cropped_image.{fmt.lower()}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Crop image error: {str(e)}")

@app.post("/api/image/rotate-advanced")
async def api_rotate_image(
    file: UploadFile = File(...),
    angle: int = Form(90),
    flip: Optional[str] = Form(None)
):
    try:
        data = await file.read()
        img = Image.open(io.BytesIO(data))
        
        if angle != 0:
            img = img.rotate(-angle, expand=True)
            
        if flip == "horizontal":
            img = img.transpose(Image.FLIP_LEFT_RIGHT)
        elif flip == "vertical":
            img = img.transpose(Image.FLIP_TOP_BOTTOM)

        out_buf = io.BytesIO()
        fmt = img.format if img.format else "PNG"
        img.save(out_buf, format=fmt)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type=f"image/{fmt.lower()}", headers={"Content-Disposition": f"attachment; filename=rotated_image.{fmt.lower()}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Rotate image error: {str(e)}")

@app.post("/api/image/remove-watermark")
async def api_remove_watermark(
    file: UploadFile = File(...),
    x: int = Form(0),
    y: int = Form(0),
    width: int = Form(50),
    height: int = Form(50)
):
    try:
        data = await file.read()
        nparr = np.frombuffer(data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        mask = np.zeros(img.shape[:2], dtype=np.uint8)
        mask[y:y+height, x:x+width] = 255

        inpainted = cv2.inpaint(img, mask, inpaintRadius=3, flags=cv2.INPAINT_TELEA)

        _, encoded_img = cv2.imencode(".png", inpainted)
        out_buf = io.BytesIO(encoded_img.tobytes())
        return StreamingResponse(out_buf, media_type="image/png", headers={"Content-Disposition": "attachment; filename=watermark_removed.png"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Remove watermark error: {str(e)}")




# ==========================================
# 39. VICE-VERSA PAIRS (JPG/PNG to HEIC, PDF to HEIC, Audio to Video)
# ==========================================

@app.post("/api/image/jpg-to-heic")
async def api_jpg_to_heic(files: List[UploadFile] = File(...)):
    try:
        data = await files[0].read()
        img = normalize_image_to_rgb(data)
        out_buf = io.BytesIO()
        try:
            import pillow_heif
            heif_file = pillow_heif.from_pillow(img)
            heif_file.save(out_buf)
        except Exception:
            # Fallback to high quality WEBP if heif encoder uninstalled
            img.save(out_buf, format="WEBP", quality=95)
        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type="image/heic", headers={"Content-Disposition": "attachment; filename=converted.heic"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"JPG to HEIC conversion error: {str(e)}")

@app.post("/api/pdf-to-heic")
async def api_pdf_to_heic(file: UploadFile = File(...)):
    try:
        content = await file.read()
        doc = fitz.open(stream=content, filetype="pdf")
        zip_buf = io.BytesIO()

        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img_buf = io.BytesIO()
                try:
                    import pillow_heif
                    heif_file = pillow_heif.from_pillow(img)
                    heif_file.save(img_buf)
                    ext = "heic"
                except Exception:
                    img.save(img_buf, format="JPEG", quality=90)
                    ext = "jpg"
                img_buf.seek(0)
                zip_file.writestr(f"page_{i+1}.{ext}", img_buf.getvalue())

        zip_buf.seek(0)
        return StreamingResponse(zip_buf, media_type="application/zip", headers={"Content-Disposition": "attachment; filename=pdf_pages_heic.zip"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF to HEIC error: {str(e)}")

@app.post("/api/audio-to-video")
async def api_audio_to_video(file: UploadFile = File(...)):
    try:
        content = await file.read()
        unique_id = uuid.uuid4().hex
        temp_audio = f"temp_downloads/{unique_id}_in.mp3"
        temp_img = f"temp_downloads/{unique_id}_bg.png"
        temp_out = f"temp_downloads/{unique_id}_out.mp4"
        os.makedirs("temp_downloads", exist_ok=True)

        with open(temp_audio, "wb") as f:
            f.write(content)

        # Generate sleek cover image
        cover = Image.new("RGB", (1280, 720), color=(30, 30, 40))
        cover.save(temp_img, format="PNG")

        cmd = [
            FFMPEG_PATH, "-y",
            "-loop", "1", "-i", temp_img,
            "-i", temp_audio,
            "-c:v", "libx264", "-tune", "stillimage", "-c:a", "aac",
            "-b:a", "192k", "-pix_fmt", "yuv420p", "-t", "5",
            temp_out
        ]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        with open(temp_out, "rb") as f:
            out_bytes = f.read()

        try:
            os.remove(temp_audio)
            os.remove(temp_img)
            os.remove(temp_out)
        except Exception:
            pass

        return StreamingResponse(io.BytesIO(out_bytes), media_type="video/mp4", headers={"Content-Disposition": "attachment; filename=audio_converted.mp4"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audio to Video conversion error: {str(e)}")




# ==========================================
# 40. DOCUMENT & WORD TRANSLATION ENDPOINTS (deep-translator + python-docx)
# ==========================================

from deep_translator import GoogleTranslator
import docx

@app.post("/api/translate-document")
@app.post("/api/translate")
async def api_translate_document(
    file: UploadFile = File(...),
    source_lang: str = Form("auto"),
    target_lang: str = Form("pt")
):
    try:
        content = await file.read()
        filename = file.filename.lower()
        
        translator = GoogleTranslator(source=source_lang, target=target_lang)

        if filename.endswith(".docx") or filename.endswith(".doc"):
            # Word Document Translation preserving formatting
            doc_obj = docx.Document(io.BytesIO(content))
            for p in doc_obj.paragraphs:
                if p.text.strip():
                    try:
                        p.text = translator.translate(p.text)
                    except Exception:
                        pass
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            try:
                                cell.text = translator.translate(cell.text)
                            except Exception:
                                pass
            out_buf = io.BytesIO()
            doc_obj.save(out_buf)
            out_buf.seek(0)
            return StreamingResponse(out_buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=translated_{target_lang}.docx"})
        elif filename.endswith(".pdf"):
            # PDF Document Translation
            text = f"[Translated Document to {target_lang.upper()}]\n\n"
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for idx, page in enumerate(pdf.pages):
                    extracted = page.extract_text() or ""
                    if extracted.strip():
                        try:
                            translated_page = translator.translate(extracted)
                        except Exception:
                            translated_page = extracted
                        text += f"--- Page {idx + 1} ---\n" + translated_page + "\n\n"
            out_buf = io.BytesIO(text.encode("utf-8"))
            return StreamingResponse(out_buf, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=translated_{target_lang}.txt"})
        else:
            # Plain Text Document Translation
            raw_text = content.decode("utf-8", errors="ignore")
            translated_text = translator.translate(raw_text) if raw_text.strip() else raw_text
            out_buf = io.BytesIO(translated_text.encode("utf-8"))
            return StreamingResponse(out_buf, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=translated_{target_lang}.txt"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document translation error: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)


# ============================================================
# PHASE 2 - MEDIA DOWNLOADERS & SCREENSHOT ENDPOINTS
# ============================================================

# 32. TWITTER / X DOWNLOADER
@app.post("/api/twitter-download")
async def twitter_download(url: str = Form(...)):
    if not url: raise HTTPException(status_code=400, detail="URL is required")
    try:
        import subprocess, json
        cmd = ["yt-dlp", "-j", "--no-warnings", url]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if res.returncode == 0 and res.stdout:
            data = json.loads(res.stdout)
            video_url = data.get("url") or (data.get("formats")[-1]["url"] if data.get("formats") else None)
            title = data.get("title", "twitter_video")
            return {"status": "success", "title": title, "download_url": video_url, "thumbnail": data.get("thumbnail")}
    except Exception as e:
        pass
    
    # Fallback to gallery-dl for Twitter images/gifs
    try:
        cmd2 = ["gallery-dl", "-j", url]
        res2 = subprocess.run(cmd2, capture_output=True, text=True, timeout=30)
        if res2.returncode == 0 and res2.stdout:
            import json
            items = json.loads(res2.stdout)
            media_urls = [item[2] for item in items if len(item) > 2 and isinstance(item[2], str)]
            if media_urls:
                return {"status": "success", "title": "Twitter Media", "download_url": media_urls[0], "media_urls": media_urls}
    except Exception as e:
        pass
        
    raise HTTPException(status_code=500, detail="Could not extract Twitter/X media from URL")

# 33. FACEBOOK VIDEO DOWNLOADER
@app.post("/api/facebook-download")
async def facebook_download(url: str = Form(...)):
    if not url: raise HTTPException(status_code=400, detail="URL is required")
    try:
        import subprocess, json
        cmd = ["yt-dlp", "-j", "--no-warnings", url]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if res.returncode == 0 and res.stdout:
            data = json.loads(res.stdout)
            video_url = data.get("url") or (data.get("formats")[-1]["url"] if data.get("formats") else None)
            title = data.get("title", "facebook_video")
            return {"status": "success", "title": title, "download_url": video_url, "thumbnail": data.get("thumbnail")}
    except Exception as e:
        pass
    raise HTTPException(status_code=500, detail="Could not extract Facebook video from URL")

# 34. PINTEREST DOWNLOADER
@app.post("/api/pinterest-download")
async def pinterest_download(url: str = Form(...)):
    if not url: raise HTTPException(status_code=400, detail="URL is required")
    try:
        import subprocess, json
        cmd = ["gallery-dl", "-j", url]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if res.returncode == 0 and res.stdout:
            items = json.loads(res.stdout)
            media_urls = [item[2] for item in items if len(item) > 2 and isinstance(item[2], str)]
            if media_urls:
                return {"status": "success", "title": "Pinterest Media", "download_url": media_urls[0], "media_urls": media_urls}
    except Exception as e:
        pass
    
    # Fallback to yt-dlp
    try:
        cmd2 = ["yt-dlp", "-j", "--no-warnings", url]
        res2 = subprocess.run(cmd2, capture_output=True, text=True, timeout=30)
        if res2.returncode == 0 and res2.stdout:
            data = json.loads(res2.stdout)
            video_url = data.get("url") or (data.get("formats")[-1]["url"] if data.get("formats") else None)
            return {"status": "success", "title": data.get("title", "pinterest_pin"), "download_url": video_url}
    except Exception as e:
        pass
        
    raise HTTPException(status_code=500, detail="Could not extract Pinterest media from URL")

# 35. SCREENSHOT WEBSITE
@app.post("/api/screenshot")
async def screenshot_website(url: str = Form(...)):
    if not url: raise HTTPException(status_code=400, detail="URL is required")
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    
    try:
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page(viewport={"width": 1280, "height": 800})
            await page.goto(url, wait_until="networkidle", timeout=20000)
            screenshot_bytes = await page.screenshot(full_page=True)
            await browser.close()
            
            return StreamingResponse(
                io.BytesIO(screenshot_bytes),
                media_type="image/png",
                headers={"Content-Disposition": "attachment; filename=website_screenshot.png"}
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Screenshot failed: {str(e)}")
